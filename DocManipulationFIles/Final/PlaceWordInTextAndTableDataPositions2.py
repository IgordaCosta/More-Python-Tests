from docx import Document

import UserImputToTableImput
import closeWord







def appendTableValue(tableTextLocations,LengthOfWordToCheck,totalLengthTableWords):
    AmountOfTimesWordApeared = len(tableTextLocations)+1

    AmountOfLettersToGoBack= LengthOfWordToCheck * AmountOfTimesWordApeared

    tableTextLocations.append(totalLengthTableWords - AmountOfLettersToGoBack)



# WordTextLocations:  [55, 79, 538, 946, 974, 1927]
# tableTextLocations:  [24, 69, 74, 80, 84, 123, 129, 153, 198, 655, 737]
# AddedAt_ 

#  _XXXXXXXXXXXXXXXXXXXXXX
# WordTextLocations:  [54, 77, 535, 942, 969, 1921]
# WordTextLocationsToAddList = ['AddedAt_54_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_77_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_535_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_942_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_969_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_1921_XXXXXXXXXXXXXXXXXXXXX']

# def GetTextAndTableDataPosition(wordDocunent,WordToCheck):
# def GetTextAndTableDataPosition(wordDocunent,WordTextLocations,tableTextLocations,AddSideSpaces=True):
# def GetTextAndTableDataPosition(wordDocument,WordTextLocations,WordTextLocationsToAddList, TableTextLocations, TableTextLocationsToAddList,AddSideSpaces=True):
def PlaceWordInTextAndTableDataPositions(wordDocument,WordYXLocations,TableYXILocation,WordTextLocations,WordTextLocationsToAddList, TableTextLocations, tableTextLocationsSimplified, UserTableImput, AddSideSpaces=True):


    # doc = open(r"C:\Users\IgorDC\Desktop\PydocTest\testFile.txt","w")#write mode 


    TableTextLocationsToAddList = UserImputToTableImput.UserImputToTableImput(tableTextLocationsSimplified, UserImput=UserTableImput)

    document = Document(wordDocument)



    ToAddValueLenSum = 0
    
    ToAddValueLen = 0

    LastItem = 'ok'
    #For ParagraphText
    for item in WordYXLocations:

        itemList = item.split(".")

        # document.paragraphs[int(itemList[0])].text[int(itemList[1])].append("TTTTTTTTTTTTTTTTTT")

        # print(document.paragraphs[int(itemList[0])].text[int(itemList[1])])

        # print(itemList)

        # print(document.paragraphs[int(itemList[0])].text)

        # print(document.paragraphs[int(itemList[0])].text[:int(itemList[1])-1])

        # print(document.paragraphs[int(itemList[0])].text[int(itemList[1])-1:])

        # print('========================================================')

        if LastItem == itemList[0]:

            ToAddValueLenSum = ToAddValueLenSum + ToAddValueLen -2
            pass
        else:
            ToAddValueLenSum = 0

            pass

        Lside = document.paragraphs[int(itemList[0])].text[:int(itemList[1])-1 + ToAddValueLenSum ]

        Rside = document.paragraphs[int(itemList[0])].text[int(itemList[1])-1 + ToAddValueLenSum:]

        ToAddValue = " "+"ÇÇÇÇÇÇÇÇÇÇÇÇÇÇÇ"

        # print(itemList)

        # print(document.paragraphs[int(itemList[0])].text)

        # print(Lside)

        # print(Rside)

        # print('========================================================')

        document.paragraphs[int(itemList[0])].text = Lside + ToAddValue + Rside

        # print(document.paragraphs[int(itemList[0])].text)

        # print("XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX")


        ToAddValueLen = len(ToAddValue)

        LastItem=itemList[0]

  


    #For Table
    for item in TableYXILocation:

        itemList = item.split(".")

        document.tables[int(itemList[0])].rows[int(itemList[1])].cells[int(itemList[2])].text = "PPPPPPPPPPPPPPPPPPPPP"


    


    


    fileName = r'C:\Users\IgorDC\Desktop\PydocTest\ChangedWordDocument20210130.docx'                 
    try:
        document.save(fileName)
    except PermissionError:
        closeWord.closeWord(fileName=fileName)
        document.save(fileName)

    # doc.close()
    print("Done")


    


WordTextLocations =  [55, 79, 538, 946, 974, 1927]
tableTextLocations =  [25, 26, 27, 28, 29, 30, 31, 32, 78, 79, 80, 81, 82, 83, 84, 90, 91, 92, 99, 100, 101, 106, 146, 147, 148, 155, 156, 157, 182, 183, 184, 185, 186, 187, 188, 189, 235, 236, 237, 238, 239, 240, 241, 699, 782]
tableTextLocationsSimplified =  [8, 7, 3, 3, 1, 3, 3, 8, 7, 1, 1]


WordTextLocationsToAddList = ['AddedAt_54_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_77_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_535_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_942_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_969_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_1921_XXXXXXXXXXXXXXXXXXXXX']


TableTextLocations =  [25, 26, 27, 28, 29, 30, 31, 32, 78, 79, 80, 81, 82, 83, 84, 90, 91, 92, 99, 100, 101, 106, 146, 147, 148, 155, 156, 157, 182, 183, 184, 185, 186, 187, 188, 189, 235, 236, 237, 238, 239, 240, 241, 699, 782]


UserTableImput = ['AddedAt_:'+str(tableTextLocationsSimplified[i])+"__:"+str(i)+"_XXXXXXXXXXXXXXXXXXXXX" for i in range(len(tableTextLocationsSimplified))]


WordToCheck= 'Ÿ'

# wordDocument = r'C:\Users\IgorDC\Desktop\PydocTest\TestDocNoChange.docx'

wordDocument = r'C:\Users\IgorDC\Desktop\PydocTest\TestDocNoChange.docx'

WordYXLocations =  ['2.25', '2.49', '2.71', '2.102', '2.128', '11.124', '23.84', '25.16', '41.96']
TableYXILocation =  ['0.0.2', '0.0.3', '0.0.4', '0.0.5', '0.0.6', '0.0.7', '0.0.8', '0.0.9', '0.1.3', '0.1.4', '0.1.5', '0.1.6', '0.1.7', 
'0.1.8', '0.1.9', '0.2.1', '0.2.2', '0.2.3', '0.2.5', '0.2.6', '0.2.7', '0.2.9', '0.3.0', '0.3.0', '0.3.0', '0.3.0', '0.3.1', '0.3.1', '0.3.1', '0.3.1', '0.3.2', '0.3.2', '0.3.2', '0.3.2', '0.3.3', '0.3.4', '0.3.5', '0.3.7', '0.3.8', '0.3.9', '1.0.2', '1.0.3', '1.0.4', '1.0.5', '1.0.6', '1.0.7', '1.0.8', '1.0.9', '1.1.3', '1.1.4', '1.1.5', '1.1.6', '1.1.7', '1.1.8', '1.1.9', '6.0.0', '6.4.4']

PlaceWordInTextAndTableDataPositions(wordDocument,WordYXLocations,TableYXILocation,WordTextLocations,WordTextLocationsToAddList, TableTextLocations, tableTextLocationsSimplified, UserTableImput, AddSideSpaces=True)



# WordTextLocations, tableTextLocations = GetTextAndTableDataPosition(wordDocunent,WordToCheck)


# print('WordTextLocations: ', WordTextLocations )

# print('tableTextLocations: ', tableTextLocations)




# results for 'Ÿ' and 'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'

# WordTextLocations:  [54, 77, 535, 942, 969, 1921]
# tableTextLocations:  [24, 69, 74, 80, 84, 123, 129, 153, 198, 655, 737]