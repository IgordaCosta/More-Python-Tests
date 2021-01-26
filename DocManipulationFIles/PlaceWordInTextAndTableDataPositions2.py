from docx import Document





def appendTableValue(tableTextLocations,LengthOfWordToCheck,totalLengthTableWords):
    AmountOfTimesWordApeared = len(tableTextLocations)+1

    AmountOfLettersToGoBack= LengthOfWordToCheck * AmountOfTimesWordApeared

    tableTextLocations.append(totalLengthTableWords - AmountOfLettersToGoBack)


WordTextLocations:  [54, 77, 535, 942, 969, 1921]

def GetTextAndTableDataPosition(wordDocunent,WordToCheck,WordTextLocations):

    document = Document(wordDocunent)

    LengthOfWordToCheck = len(WordToCheck)

    #to get text letter locations

    WordTextLocations=[]
    totalLengthWords=0

    WordTextCount=0
    for paragraph in document.paragraphs:
        for letter in paragraph.text:
        
            totalLengthWords = totalLengthWords + 1 
            if WordTextCount == letter:

                AmountOfTimesWordApeared = len(WordTextLocations)+1

                AmountOfLettersToGoBack= LengthOfWordToCheck * AmountOfTimesWordApeared

                WordTextLocations.append(totalLengthWords + totalLengthWords - AmountOfLettersToGoBack)



    # To get table data letter locations

    lastLetter = ''
    tableTextLocations = []
    totalLengthTableWords=0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:

                CellText=cell.text
                # print(CellText)
                
                for letter in CellText:
                
                    # print(word.text)
                    LengthOfWordToCheck = 1
                    totalLengthTableWords = totalLengthTableWords + LengthOfWordToCheck
                        
                    if WordToCheck == letter:
                        if letter == lastLetter:
                            lastLetter = ''
                            totalLengthTableWords = totalLengthTableWords - LengthOfWordToCheck
                        else:
                            print('letter: ',letter)

                            appendTableValue(tableTextLocations,LengthOfWordToCheck,totalLengthTableWords)

                    lastLetter = letter

                        
    # print(len(WordTextLocations)+len(tableTextLocations))

    return WordTextLocations, tableTextLocations







WordToCheck= 'Ÿ'

wordDocunent = r'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'




# WordTextLocations, tableTextLocations = GetTextAndTableDataPosition(wordDocunent,WordToCheck)


# print('WordTextLocations: ', WordTextLocations )

# print('tableTextLocations: ', tableTextLocations)




# results for 'Ÿ' and 'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'

# WordTextLocations:  [55, 79, 538, 946, 974, 1927]
# tableTextLocations:  [24, 69, 74, 80, 84, 123, 129, 153, 198, 655, 737]