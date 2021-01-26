from docx import Document





# def appendTableValue(tableTextLocations,LengthOfWordToCheck,totalLengthTableWords):
#     # AmountOfTimesWordApeared = len(tableTextLocations)+1

#     # AmountOfLettersToGoBack= LengthOfWordToCheck * AmountOfTimesWordApeared

#     # tableTextLocations.append(totalLengthTableWords - AmountOfLettersToGoBack)
#     tableTextLocations.append(totalLengthTableWords)

def appendTableValueSimlified(tableTextLocationsSimplified,LengthOfWordToCheck,totalLengthTableWords):
    AmountOfTimesWordApeared = len(tableTextLocationsSimplified)+1

    AmountOfLettersToGoBack= LengthOfWordToCheck * AmountOfTimesWordApeared

    tableTextLocationsSimplified.append(totalLengthTableWords - AmountOfLettersToGoBack)




def GetTextAndTableDataPosition(wordDocunent,WordToCheck):

    document = Document(wordDocunent)

    LengthOfWordToCheck = len(WordToCheck)

    #to get text letter locations

    file1 = open(r"C:\Users\IgorDC\Desktop\PydocTest\testFile.txt","w")#write mode 
    


    WordTextLocations=[]
    totalLengthWords=0
    for paragraph in document.paragraphs:
        for letter in paragraph.text:

            
        
            totalLengthWords = totalLengthWords + 1 
            file1.write('letter: '+str(letter)+ '  totalLengthWords:  '+str(totalLengthWords) + "\n") 
            # print(letter, totalLengthWords)
            if WordToCheck == letter:

                AmountOfTimesWordApeared = len(WordTextLocations)+1

                AmountOfLettersToGoBack= LengthOfWordToCheck * AmountOfTimesWordApeared

                # WordTextLocations.append(totalLengthWords - AmountOfLettersToGoBack)

                WordTextLocations.append(totalLengthWords)


    file1.close() 
    # To get table data letter locations

    lastLetter = ''
    tableTextLocations = []
    tableTextLocationsSimplified = []
    totalLengthTableWords=0
    totalLengthTableWordsOld=0
    Cellsize = 1
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:

                CellText=cell.text
                print(CellText)
                
                # step = 0
                for letter in CellText:
                
                    # print(word.text)
                    LengthOfWordToCheck = 1
                    totalLengthTableWords = totalLengthTableWords + LengthOfWordToCheck

                    step = 0
                    
                    if WordToCheck == letter:
                        # if letter == lastLetter:
                        #     lastLetter = ''
                        #     totalLengthTableWords = totalLengthTableWords - LengthOfWordToCheck
                        # else:
                        #     # print('letter: ',letter)
                        #     appendTableValueSimlified(tableTextLocationsSimplified,LengthOfWordToCheck,totalLengthTableWords)
                        
                       
                        sizeOftableTextLocations = len(tableTextLocations)

                        if sizeOftableTextLocations == 0:
                            
                            totalLengthTableWordsOld = totalLengthTableWords

                            tableTextLocations.append(totalLengthTableWords)

                        else:

                            oldtableTextLocationsValue = tableTextLocations[sizeOftableTextLocations - 1]
                            

                            tableTextLocations.append(totalLengthTableWords)

                            if oldtableTextLocationsValue + 1 == totalLengthTableWords:
                                Cellsize = Cellsize + 1
                            else:
                                tableTextLocationsSimplified.append(Cellsize)
                                Cellsize = 1

                            totalLengthTableWordsOld = totalLengthTableWords


                    # step = step + 1


                    # lastLetter = letter

                        
    # print(len(WordTextLocations)+len(tableTextLocations))

    return WordTextLocations, tableTextLocations, tableTextLocationsSimplified







WordToCheck= 'Ÿ'

wordDocunent = r'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'




WordTextLocations, tableTextLocations, tableTextLocationsSimplified = GetTextAndTableDataPosition(wordDocunent,WordToCheck)


print('WordTextLocations: ', WordTextLocations )

print('tableTextLocations: ', tableTextLocations)

print("tableTextLocationsSimplified: ", tableTextLocationsSimplified)




# results for 'Ÿ' and 'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'

# WordTextLocations:  [55, 79, 538, 946, 974, 1927]
# tableTextLocations:  [24, 69, 74, 80, 84, 123, 129, 153, 198, 655, 737]

# tableTextLocations:  [25, 26, 27, 28, 29, 30, 31, 32, 78, 79, 80, 81, 82, 83, 84, 90, 91, 92, 99, 100, 101, 106, 146, 147, 148, 155, 156, 157, 182, 183, 184, 185, 186, 187, 188, 189, 235, 236, 237, 238, 239, 240, 241, 699, 782]