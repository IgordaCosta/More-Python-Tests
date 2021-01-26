from docx import Document







def appendTableValue(tableTextLocations,LengthOfWordToCheck,totalLengthTableWords):
    AmountOfTimesWordApeared = len(tableTextLocations)+1

    AmountOfLettersToGoBack= LengthOfWordToCheck * AmountOfTimesWordApeared

    tableTextLocations.append(totalLengthTableWords - AmountOfLettersToGoBack)



# WordTextLocations:  [55, 79, 538, 946, 974, 1927]
# tableTextLocations:  [24, 69, 74, 80, 84, 123, 129, 153, 198, 655, 737]
# AddedAt_ 

#  _XXXXXXXXXXXXXXXXXXXXXX
# WordTextLocations:  [54, 77, 535, 942, 969, 1921]
# WordTextLocationsToAddList = ['AddedAt_54_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_77_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_535_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_942_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_969_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_1921_XXXXXXXXXXXXXXXXXXXXX']

# def GetTextAndTableDataPosition(wordDocunent,WordToCheck):
# def GetTextAndTableDataPosition(wordDocunent,WordTextLocations,tableTextLocations,AddSideSpaces=True):
def GetTextAndTableDataPosition(wordDocument,WordTextLocations,WordTextLocationsToAddList, TableTextLocations, TableTextLocationsToAddList,AddSideSpaces=True):

    document = Document(wordDocument)

    AtWordStep = -1

    DocumentLegth=0
    for i in range(len(document.paragraphs)):
       
        paragraphLength=len(document.paragraphs[i].text)

        DocumentLegthOld= DocumentLegth

        DocumentLegth = DocumentLegth + paragraphLength

        AllLenWord = 0
        insideParagraph = True
        numberOfRunsInside = 0
        while insideParagraph:
            paragraphLength=len(document.paragraphs[i].text)

            try:
                AtWordLocation = WordTextLocations[AtWordStep +1]


                if DocumentLegth >= AtWordLocation:
                    if  AtWordLocation >= DocumentLegthOld:

                        AtWordStep = AtWordStep + 1

                        LocationToAddData=AtWordLocation-DocumentLegthOld

                        TotalSpaceToForward = LocationToAddData 

                        WordToAdd=WordTextLocationsToAddList[AtWordStep]

                        if AddSideSpaces:
                            WordToAdd= " "+ WordToAdd + " "
                        
                        document.paragraphs[i].text = document.paragraphs[i].text[:TotalSpaceToForward + AllLenWord -1 ] +WordToAdd + document.paragraphs[i].text[TotalSpaceToForward + AllLenWord: ]

                        LenWord= len(WordToAdd)

                        AllLenWord= AllLenWord + LenWord -1

                        numberOfRunsInside= numberOfRunsInside + 1
                    
                    else:
                        insideParagraph = False
                
                else:
                    insideParagraph = False
                

            except IndexError:
                insideParagraph = False



        # print(document.paragraphs[i].text)
        
        

    





        # if WordToCheck in paragraph.text:
        # print(paragraph.text)
        # for letter in paragraph.text:

        

        #     if 
            # print(len(paragraph.text))
        
            # totalLengthWords = totalLengthWords + 1 

            # print(totalLengthWords)

            # print(WordTextLocations)

            

        #     if totalLengthWords == CurrentSizeIndexLocations:

        #         if AddSideSpaces:
        #             letter = ' ' + WordTextLocationsToAddList[CurrentIndex] + ' '
        #         else:
        #             letter = WordTextLocationsToAddList[CurrentIndex]

        #         # print(letter)

        #         lengthOfWordAdded = len(WordTextLocationsToAddList[CurrentIndex])
        #         totalLengthWords= totalLengthWords + lengthOfWordAdded

        #         CurrentIndex=CurrentIndex+1

        #         try:
        #             CurrentSizeIndexLocations = CurrentSizeIndexLocations + WordTextLocations[CurrentIndex]
        #         except IndexError:
        #             pass

        #     paragraph.text = letter
        # print(paragraph.text)

    
            

            # if WordToCheck == letter:

            #     AmountOfTimesWordApeared = len(WordTextLocations)+1

            #     AmountOfLettersToGoBack= LengthOfWordToCheck * AmountOfTimesWordApeared

            #     WordTextLocations.append(totalLengthWords - AmountOfLettersToGoBack)



    #To get table data letter locations

    lastLetter = ''
    tableTextLocations = []
    totalLengthTableWords=0

    AtWordStep = -1

    DocumentLegth=0
    for table in document.tables:
        for row in table.rows:
            # for cell in row.cells:

                # CellText=cell.text
            
            for i in range(len(row.cells)):

                CellText=row.cells[i].text

                # print(CellText)


                paragraphLength=len(CellText)

                DocumentLegthOld= DocumentLegth

                DocumentLegth = DocumentLegth + paragraphLength

                # print(DocumentLegth)

                AllLenWord = 0
                insideParagraph = True
                numberOfRunsInside = 0
                print("Item: " +str(DocumentLegth))
                print("CellText Before: ", CellText)
                while insideParagraph:
                    paragraphLength=len(CellText)

                    

                    try:
                        AtWordLocation = TableTextLocations[AtWordStep +1]

                        # print('DocumentLegthOld: ',DocumentLegthOld)
                        # print('DocumentLegth: ',DocumentLegth)
                        # print('AtWordLocation: ',AtWordLocation)
                        # print("=====================================================")


                        if DocumentLegth >= AtWordLocation:
                            if  AtWordLocation >= DocumentLegthOld:

                                AtWordStep = AtWordStep + 1

                                LocationToAddData=AtWordLocation-DocumentLegthOld

                                TotalSpaceToForward = LocationToAddData 

                                WordToAdd=TableTextLocationsToAddList[AtWordStep]

                                if AddSideSpaces:
                                    WordToAdd= " "+ WordToAdd + " "
                                
                                CellText = CellText[:TotalSpaceToForward + AllLenWord -1 ] +WordToAdd + CellText[TotalSpaceToForward + AllLenWord: ]

                                LenWord= len(WordToAdd)

                                AllLenWord= AllLenWord + LenWord -1

                                numberOfRunsInside= numberOfRunsInside + 1

                                print("CellText: ", CellText)
                            
                            else:
                                insideParagraph = False
                        
                        else:
                            insideParagraph = False
                        

                    except IndexError:
                        insideParagraph = False
                


                print("CellText After: ", CellText)
                print("=================================================")





                # for letter in CellText:
                #     print(letter)
                
                #     # print(word.text)
                #     LengthOfWordToCheck = 1
                #     totalLengthTableWords = totalLengthTableWords + LengthOfWordToCheck
                        
                #     if WordToCheck == letter:
                #         if letter == lastLetter:
                #             lastLetter = ''
                #             totalLengthTableWords = totalLengthTableWords - LengthOfWordToCheck
                #         else:
                #             # print('letter: ',letter)
                #             appendTableValue(tableTextLocations,LengthOfWordToCheck,totalLengthTableWords)

                #     lastLetter = letter

                        
    # print(len(WordTextLocations)+len(tableTextLocations))

    # return WordTextLocations, tableTextLocations




WordTextLocations =  [55, 79, 538, 946, 974, 1927]

WordTextLocationsToAddList = ['AddedAt_54_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_77_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_535_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_942_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_969_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_1921_XXXXXXXXXXXXXXXXXXXXX']

# TableTextLocations = [24, 69, 74, 80, 84, 123, 129, 153, 198, 655, 737]



TableTextLocations =  [25, 26, 27, 28, 29, 30, 31, 32, 78, 79, 80, 81, 82, 83, 84, 90, 91, 92, 99, 100, 101, 106, 146, 147, 148, 155, 156, 157, 182, 183, 184, 185, 186, 187, 188, 189, 235, 236, 237, 238, 239, 240, 241, 699, 782]

# TableTextLocationsToAddList = ['AddedAt_24', 'AddedAt_69', 'AddedAt_74', 'AddedAt_80', 'AddedAt_84', 'AddedAt_123', 'AddedAt_129', 'AddedAt_153', 'AddedAt_198', 'AddedAt_655', 'AddedAt_737']

# TableTextLocationsToAddList = ['25', '26', '27', '28', '29', '30', '31', '32', '78', '79', '80', '81', '82', '83', '84', '90', '91', '92', '99', '100', '101', '106', '146', '147', '148', '155', '156', '157', '182', '183', '184', '185', '186', '187', '188', '189', '235', '236', '237', '238', '239', '240', '241', '699', '782']

TableTextLocationsToAddList = ['AddedAt_'+str(i)+"_XXXXXXXXXXXXXXXXXXXXX" for i in TableTextLocations]

# print(TableTextLocationsToAddList)



WordToCheck= 'Ÿ'

wordDocument = r'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'

GetTextAndTableDataPosition(wordDocument,WordTextLocations,WordTextLocationsToAddList, TableTextLocations, TableTextLocationsToAddList,AddSideSpaces=True)




# WordTextLocations, tableTextLocations = GetTextAndTableDataPosition(wordDocunent,WordToCheck)


# print('WordTextLocations: ', WordTextLocations )

# print('tableTextLocations: ', tableTextLocations)




# results for 'Ÿ' and 'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'

# WordTextLocations:  [54, 77, 535, 942, 969, 1921]
# tableTextLocations:  [24, 69, 74, 80, 84, 123, 129, 153, 198, 655, 737]